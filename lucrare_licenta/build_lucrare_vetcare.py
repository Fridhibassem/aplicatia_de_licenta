from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BASE_DIR = Path(__file__).resolve().parents[1]
WORK_DIR = Path(__file__).resolve().parent
OUT = "Lucrare_licenta_VetCare_minim_40_pagini.docx"


def set_cell_text(cell, text, bold=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, before, after in [
        ("Heading 1", 16, 18, 10),
        ("Heading 2", 14, 12, 6),
        ("Heading 3", 12, 8, 4),
    ]:
        st = styles[style_name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.5


def add_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("VetCare - Lucrare de licență | Pagina ")
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    add_page_field(p)


def add_center(doc, text, size=12, bold=False, after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    r.font.size = Pt(size)
    return p


def add_para(doc, text, style=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    r.font.size = Pt(12)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    r.font.size = Pt(12)


def add_table(doc, headers, rows, widths=None, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, size=font_size)
        shade(hdr[i], "EDEDED")
        if widths:
            set_width(hdr[i], widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val), size=font_size)
            if widths:
                set_width(cells[i], widths[i])
    doc.add_paragraph()
    return table


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.right_indent = Cm(0.2)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    for line in code.strip().splitlines():
        r = p.add_run(line + "\n")
        r.font.name = "Courier New"
        r._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
        r._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
        r.font.size = Pt(9)


def add_figure(doc, image_path, caption, width_cm=15.5):
    path = Path(image_path)
    if not path.exists():
        add_para(doc, f"[Captură lipsă: {path.name}] {caption}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption)
    r.italic = True
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    r.font.size = Pt(10)


def make_code_image(source_path, output_path, title, max_lines=42):
    source = Path(source_path)
    output = Path(output_path)
    text = source.read_text(encoding="utf-8", errors="replace").splitlines()
    lines = text[:max_lines]
    try:
        font = ImageFont.truetype("/System/Library/Fonts/Menlo.ttc", 18)
        title_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 22)
    except Exception:
        font = ImageFont.load_default()
        title_font = font
    line_h = 25
    w = 1500
    h = 70 + line_h * (len(lines) + 1)
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    d.rectangle([0, 0, w, 52], fill=(31, 58, 95))
    d.text((24, 14), title, font=title_font, fill="white")
    y = 68
    for i, line in enumerate(lines, start=1):
        if i % 2 == 0:
            d.rectangle([0, y - 3, w, y + line_h - 3], fill=(246, 248, 250))
        d.text((22, y), f"{i:02d}", font=font, fill=(120, 120, 120))
        d.text((78, y), line[:145], font=font, fill=(20, 20, 20))
        y += line_h
    output.parent.mkdir(parents=True, exist_ok=True)
    img.save(output)
    return output


def cover(doc, title_page=False):
    add_center(doc, "MINISTERUL EDUCAȚIEI ȘI CERCETĂRII", 12, True, 2)
    add_center(doc, "UNIVERSITATEA „OVIDIUS” DIN CONSTANȚA", 12, True, 2)
    add_center(doc, "FACULTATEA DE MATEMATICĂ ȘI INFORMATICĂ", 12, True, 2)
    add_center(doc, "SPECIALIZAREA INFORMATICĂ", 12, True, 72)
    add_center(doc, "LUCRARE DE LICENȚĂ", 18, True, 20)
    if title_page:
        add_center(doc, "Dezvoltarea unei aplicații web pentru gestionarea unui cabinet veterinar", 16, True, 2)
        add_center(doc, "Studiu de caz: platforma VetCare", 16, True, 70)
    else:
        add_center(doc, " ", 16, True, 88)
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_text(table.cell(0, 0), "Coordonator științific:", True, 12)
    set_cell_text(table.cell(0, 1), "[Titlu, Nume coordonator]", False, 12)
    set_cell_text(table.cell(1, 0), "Absolvent:", True, 12)
    set_cell_text(table.cell(1, 1), "[Numele studentului]", False, 12)
    for row in table.rows:
        for cell in row.cells:
            set_width(cell, 7.2)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                elem = OxmlElement(f"w:{edge}")
                elem.set(qn("w:val"), "nil")
                tc_borders.append(elem)
            tc_pr.append(tc_borders)
    add_center(doc, "", 12, False, 110)
    add_center(doc, "CONSTANȚA", 12, True, 0)
    add_center(doc, "2026", 12, True, 0)


def abstracts(doc):
    add_heading(doc, "Rezumat", 1)
    for text in [
        "Lucrarea prezintă proiectarea și implementarea unei aplicații web full-stack pentru gestionarea unui cabinet veterinar, denumită VetCare. Aplicația oferă funcții publice de prezentare a cabinetului, formular de contact, autentificare, programare online, profil utilizator și panou administrativ. Tema este motivată de nevoia digitalizării serviciilor locale de îngrijire veterinară, în care comunicarea rapidă, evidența programărilor și accesul controlat la date au impact direct asupra calității relației dintre cabinet și clienți.",
        "Soluția este dezvoltată cu Angular 19 pe partea de frontend, Node.js și Express pe partea de backend și MySQL pentru persistența datelor. Sistemul folosește JSON Web Tokens pentru autentificare, bcrypt pentru protejarea parolelor, validări server-side, rate limiting pentru endpointurile sensibile și constrângeri de bază de date pentru păstrarea integrității datelor. Contribuția proprie constă în structurarea funcțională a aplicației, implementarea fluxului de programare cu verificarea sloturilor ocupate, separarea rolurilor utilizator/admin și realizarea unei documentații tehnice coerente.",
        "Cuvinte cheie: aplicație web, cabinet veterinar, Angular, Node.js, Express, MySQL, JWT, securitate web, programări online.",
    ]:
        add_para(doc, text)
    add_heading(doc, "Abstract", 1)
    for text in [
        "This thesis presents the design and implementation of a full-stack web application for managing a veterinary clinic, named VetCare. The application provides public presentation pages, a contact form, authentication, online appointment booking, user profile management and an administrative dashboard. The topic is motivated by the need to digitize local veterinary services, where fast communication, appointment traceability and controlled access to data improve the relationship between the clinic and its clients.",
        "The solution is implemented using Angular 19 for the frontend, Node.js and Express for the backend, and MySQL for data persistence. The system uses JSON Web Tokens for authentication, bcrypt for password protection, server-side validation, rate limiting for sensitive endpoints and relational constraints for data integrity. The original contribution consists of the functional design of the system, the appointment workflow with occupied-slot validation, role-based separation between regular users and administrators, and a coherent technical documentation of the software product.",
        "Keywords: web application, veterinary clinic, Angular, Node.js, Express, MySQL, JWT, web security, online appointments.",
    ]:
        add_para(doc, text)


def front_matter(doc):
    add_heading(doc, "Cuprins", 1)
    contents = [
        ("1. Introducere", "1"),
        ("2. Analiza domeniului și definirea cerințelor", "5"),
        ("3. Tehnologii utilizate", "11"),
        ("4. Arhitectura și proiectarea sistemului", "17"),
        ("5. Proiectarea bazei de date și a API-ului", "24"),
        ("6. Implementarea aplicației VetCare", "31"),
        ("7. Testare, validare și evaluare", "38"),
        ("8. Concluzii și direcții viitoare", "45"),
        ("Bibliografie", "48"),
        ("Anexe", "50"),
    ]
    add_table(doc, ["Secțiune", "Pagina"], contents, [13.0, 2.2], 10)
    add_para(doc, "Notă: numerotarea paginilor este orientativă în această versiune draft. În Word, cuprinsul poate fi actualizat după completarea numelui studentului, a coordonatorului și a eventualelor capturi de ecran finale.")
    add_heading(doc, "Lista figurilor", 1)
    add_table(doc, ["Nr.", "Descriere"], [
        ("Figura 1", "Arhitectura generală a aplicației VetCare"),
        ("Figura 2", "Fluxul de autentificare cu JWT"),
        ("Figura 3", "Modelul relațional al bazei de date"),
        ("Figura 4", "Fluxul programării online"),
        ("Figura 5", "Structura componentelor Angular"),
    ], [2.2, 13.0], 10)
    add_heading(doc, "Lista tabelelor", 1)
    add_table(doc, ["Nr.", "Descriere"], [
        ("Tabelul 1", "Cerințe funcționale"),
        ("Tabelul 2", "Cerințe nefuncționale"),
        ("Tabelul 3", "Tehnologii utilizate"),
        ("Tabelul 4", "Structura tabelelor MySQL"),
        ("Tabelul 5", "Endpointuri API"),
        ("Tabelul 6", "Scenarii de test"),
    ], [2.2, 13.0], 10)


def introduction(doc):
    add_heading(doc, "1. Introducere", 1)
    paras = [
        "Digitalizarea serviciilor medicale și conexe a devenit o cerință firească pentru organizațiile care lucrează direct cu publicul. În cazul cabinetelor veterinare, această transformare este vizibilă prin nevoia de programări online, comunicare rapidă, evidență a pacienților și reducerea timpului petrecut la telefon pentru operații administrative simple. Un cabinet veterinar modern nu mai funcționează doar ca spațiu fizic, ci și ca serviciu digital disponibil pentru clienți în afara programului clasic.",
        "Tema lucrării se încadrează în zona aplicațiilor web full-stack și urmărește realizarea unei platforme pentru gestionarea relației dintre un cabinet veterinar și clienții săi. Aplicația VetCare oferă o interfață publică prin care vizitatorii pot afla informații despre cabinet, pot trimite mesaje și pot crea conturi, dar include și componente protejate pentru programări, profil utilizator și administrare. Din punct de vedere informatic, proiectul combină noțiuni de analiză a cerințelor, proiectare de baze de date, programare frontend, dezvoltare backend, securitate și testare.",
        "Motivația personală pentru alegerea temei este legată de caracterul practic al aplicației. Spre deosebire de un proiect strict demonstrativ, VetCare pornește de la un scenariu clar: un client dorește să programeze un animal la consultație, iar administratorul cabinetului trebuie să urmărească rapid programările, utilizatorii și mesajele primite. Această situație obligă aplicația să trateze probleme reale, cum ar fi autentificarea, evitarea conflictelor de orar, validarea datelor introduse și diferențierea drepturilor de acces.",
        "Lucrarea are ca obiectiv principal proiectarea și implementarea unei aplicații web care să fie ușor de folosit, coerentă la nivel de arhitectură și suficient de sigură pentru un context educațional. Obiectivele secundare sunt definirea cerințelor funcționale și nefuncționale, alegerea unei stive tehnologice potrivite, proiectarea bazei de date, implementarea API-ului REST, realizarea interfeței Angular și descrierea unui plan de testare care să acopere fluxurile importante.",
        "Din punct de vedere istoric, aplicațiile web au evoluat de la pagini statice la sisteme complexe în care clientul și serverul colaborează permanent. Separarea dintre frontend și backend permite dezvoltarea independentă a interfeței și a logicii de business. VetCare folosește această separare printr-un frontend Angular care comunică prin HTTP cu un backend Express, iar backend-ul păstrează datele într-o bază MySQL. Această arhitectură este potrivită pentru aplicații educaționale, dar și pentru produse mici sau medii care pot fi extinse ulterior.",
        "Contribuția proprie constă în definirea fluxului complet al aplicației, implementarea componentelor principale și documentarea deciziilor tehnice. Aplicația nu este prezentată ca un sistem clinic complet, ci ca un nucleu funcțional care poate fi extins cu evidența animalelor, fișe medicale, facturare sau notificări. În forma actuală, accentul este pus pe programări, comunicare și administrare, acestea fiind funcțiile care demonstrează cel mai bine legătura dintre interfață, API și baza de date.",
    ]
    for p in paras:
        add_para(doc, p)
    add_heading(doc, "1.1 Structura lucrării", 2)
    for p in [
        "Capitolul al doilea prezintă analiza domeniului și cerințele aplicației. Sunt discutate actorii principali, cazurile de utilizare și constrângerile de securitate și utilizabilitate. Capitolul al treilea descrie tehnologiile folosite și justifică alegerea lor în raport cu obiectivele proiectului.",
        "Capitolul al patrulea descrie arhitectura generală, fluxurile aplicației și organizarea codului. Capitolul al cincilea se concentrează pe baza de date și pe API-ul backend. Capitolul al șaselea prezintă implementarea concretă, iar capitolul al șaptelea discută testarea. Ultimul capitol sintetizează rezultatele și propune direcții viitoare de dezvoltare.",
    ]:
        add_para(doc, p)


def requirements(doc):
    add_heading(doc, "2. Analiza domeniului și definirea cerințelor", 1)
    for p in [
        "Un cabinet veterinar gestionează simultan servicii medicale, comunicare cu proprietarii animalelor și activități administrative. Într-un flux tradițional, multe activități sunt realizate telefonic sau prin registre separate, ceea ce poate genera întârzieri, dublări ale informațiilor și dificultăți în urmărirea istoricului. O aplicație web nu elimină rolul personalului, dar poate automatiza acțiunile repetitive și poate centraliza datele necesare pentru decizii rapide.",
        "Analiza domeniului a pornit de la identificarea actorilor: vizitatorul neautentificat, utilizatorul autentificat și administratorul. Vizitatorul poate consulta informații publice și poate trimite un mesaj prin formularul de contact. Utilizatorul autentificat poate face programări și poate vedea istoricul propriu. Administratorul are acces la dashboard, statistici, lista utilizatorilor, programări și mesaje.",
        "O cerință importantă este separarea clară a drepturilor. Un utilizator obișnuit nu trebuie să poată vedea programările altor persoane sau să schimbe roluri. În același timp, administratorul trebuie să poată gestiona datele operaționale ale cabinetului. Această diferențiere este implementată prin tokenuri JWT care includ rolul utilizatorului și prin middleware backend care verifică accesul înainte de executarea operațiilor sensibile.",
    ]:
        add_para(doc, p)
    add_heading(doc, "2.1 Cerințe funcționale", 2)
    functional_rows = [
        ("CF1", "Înregistrare cont", "Aplicația permite crearea unui cont folosind email, parolă și nume."),
        ("CF2", "Autentificare", "Utilizatorul se poate autentifica și primește un token JWT."),
        ("CF3", "Profil utilizator", "Utilizatorul poate vizualiza datele proprii și poate modifica numele sau parola."),
        ("CF4", "Programare online", "Utilizatorul alege data, ora, serviciul și tipul animalului."),
        ("CF5", "Verificare sloturi", "Aplicația afișează orele ocupate pentru data selectată."),
        ("CF6", "Anulare programare", "Utilizatorul poate anula programări proprii, iar administratorul poate gestiona toate programările."),
        ("CF7", "Formular contact", "Vizitatorii pot trimite mesaje fără autentificare."),
        ("CF8", "Dashboard admin", "Administratorul vede statistici despre utilizatori, programări și mesaje."),
        ("CF9", "Gestionare utilizatori", "Administratorul poate schimba roluri și poate șterge conturi."),
        ("CF10", "Gestionare mesaje", "Administratorul poate marca mesajele ca citite și le poate șterge."),
    ]
    add_table(doc, ["Cod", "Cerință", "Descriere"], functional_rows, [2.0, 4.0, 9.2], 9)
    for p in [
        "Cerințele funcționale au fost alese astfel încât aplicația să acopere ciclul minim al interacțiunii dintre client și cabinet. Accentul cade pe programare și comunicare, deoarece acestea sunt procesele care se repetă frecvent și care pot fi simplificate printr-o platformă web.",
        "Fluxul programării are o importanță specială. Sistemul trebuie să împiedice alegerea zilelor din trecut, a weekendurilor și a intervalelor în afara programului. De asemenea, același slot nu trebuie să fie acceptat de două ori. Această logică este verificată la nivel de backend, deoarece interfața poate fi ocolită prin cereri HTTP directe.",
    ]:
        add_para(doc, p)
    add_heading(doc, "2.2 Cerințe nefuncționale", 2)
    nonfunctional_rows = [
        ("CNF1", "Securitate", "Parolele sunt hash-uite, iar accesul la rute protejate se face cu JWT."),
        ("CNF2", "Utilizabilitate", "Interfața este în limba română și folosește mesaje clare de eroare."),
        ("CNF3", "Performanță", "API-ul trebuie să răspundă rapid pentru operații curente de listare și creare."),
        ("CNF4", "Mentenabilitate", "Codul este separat în componente, rute, controllere, servicii și middleware."),
        ("CNF5", "Integritate", "Baza de date folosește chei primare, chei externe și indecși."),
        ("CNF6", "Portabilitate", "Proiectul poate fi rulat local cu Node.js, npm și MySQL."),
    ]
    add_table(doc, ["Cod", "Categorie", "Descriere"], nonfunctional_rows, [2.0, 4.0, 9.2], 9)
    for p in [
        "Cerințele nefuncționale influențează calitatea soluției chiar dacă nu sunt observate imediat de utilizator. De exemplu, o aplicație poate permite autentificarea, dar dacă parolele sunt stocate în clar, soluția este inacceptabilă din punct de vedere al securității. În VetCare, parolele sunt protejate cu bcrypt, iar endpointurile de autentificare sunt limitate pentru a reduce riscul încercărilor repetate.",
        "Mentenabilitatea a fost urmărită prin organizarea proiectului în două aplicații distincte. Frontend-ul conține componente Angular specializate pe pagini și servicii pentru comunicarea cu API-ul. Backend-ul conține rute, controllere, middleware și un adaptor pentru baza de date. Această împărțire facilitează modificările ulterioare, deoarece o schimbare în interfață nu obligă rescrierea logicii de acces la date.",
    ]:
        add_para(doc, p)
    add_heading(doc, "2.3 Cazuri de utilizare", 2)
    use_cases = [
        "Crearea contului: utilizatorul completează email, nume și parolă; backend-ul validează datele, verifică unicitatea emailului, hash-uiește parola și returnează tokenul.",
        "Programarea unei consultații: utilizatorul autentificat deschide pagina de programare, alege o zi lucrătoare, selectează un interval liber, completează serviciul și tipul animalului, apoi trimite cererea.",
        "Administrarea programărilor: administratorul accesează dashboard-ul, vede programările ordonate, filtrează mental după status și poate confirma, anula sau șterge înregistrări.",
        "Trimiterea unui mesaj de contact: vizitatorul completează formularul public, iar mesajul este stocat în tabela contact_messages pentru a fi consultat ulterior de administrator.",
    ]
    for uc in use_cases:
        add_bullet(doc, uc)
    add_para(doc, "Aceste cazuri de utilizare acoperă atât partea publică, cât și partea protejată a aplicației. Ele sunt suficiente pentru demonstrarea unei aplicații full-stack complete, cu fluxuri de citire, scriere, actualizare și ștergere a datelor.")


def technologies(doc):
    add_heading(doc, "3. Tehnologii utilizate", 1)
    for p in [
        "Alegerea tehnologiilor a urmărit obținerea unui echilibru între claritate, productivitate și relevanță pentru dezvoltarea web modernă. Proiectul folosește TypeScript și Angular pentru interfață, Node.js și Express pentru API, MySQL pentru baza de date și biblioteci consacrate pentru autentificare și securitate.",
        "Angular este potrivit pentru aplicații cu mai multe pagini și componente reutilizabile. În VetCare, fiecare pagină importantă este reprezentată de un component standalone: Home, About, Contact, Login, Programare, Profile și Dashboard. Rutarea este definită într-un fișier central, iar protecția paginilor se face cu guard-uri.",
        "Node.js permite rularea JavaScript pe server și este adecvat pentru API-uri HTTP. Express adaugă un strat simplu pentru definirea rutelor și middleware-urilor. În proiect, Express este folosit pentru rutele de autentificare, programări, dashboard și contact.",
        "MySQL a fost ales deoarece modelul aplicației este relațional: utilizatorii au programări, iar mesajele de contact sunt stocate separat. Tabelele folosesc chei primare UUID, indecși și o relație între users și appointments. Această structură este ușor de explicat, testat și extins.",
    ]:
        add_para(doc, p)
    rows = [
        ("Angular 19", "Frontend", "Componente standalone, rutare, formulare și comunicare HTTP."),
        ("TypeScript 5.7", "Limbaj frontend", "Tipare statică și cod mai ușor de întreținut."),
        ("Node.js", "Runtime backend", "Execuția serverului JavaScript și a API-ului."),
        ("Express", "Framework backend", "Rute HTTP, middleware și gestionarea erorilor."),
        ("MySQL", "Bază de date", "Persistența utilizatorilor, programărilor și mesajelor."),
        ("JWT", "Autentificare", "Transmiterea identității și a rolului între client și server."),
        ("bcryptjs", "Securitate parole", "Hash-uirea parolelor înainte de stocare."),
        ("express-rate-limit", "Protecție API", "Limitarea încercărilor repetate pe rute sensibile."),
    ]
    add_table(doc, ["Tehnologie", "Rol", "Utilizare în proiect"], rows, [3.2, 3.2, 8.8], 9)
    add_heading(doc, "3.1 Definiții ale tehnologiilor și conceptelor", 2)
    definition_rows = [
        ("HTML", "Limbaj de marcare folosit pentru structura paginilor web. În VetCare apare în template-urile componentelor Angular, de exemplu home.component.html, login.component.html și programare.component.html."),
        ("CSS", "Limbaj de stilizare folosit pentru aspect vizual, layout, culori, spațiere și responsive design. În proiect, fiecare componentă are fișier CSS separat."),
        ("TypeScript", "Limbaj bazat pe JavaScript, cu tipare statică. În Angular este folosit pentru logica componentelor, servicii, guard-uri și configurarea rutelor."),
        ("JWT", "JSON Web Token este un format compact pentru transmiterea informațiilor semnate între client și server. În VetCare conține id-ul, emailul și rolul utilizatorului."),
        ("API", "Application Programming Interface reprezintă interfața prin care două aplicații comunică. În VetCare, API-ul backend oferă rute pentru autentificare, programări, mesaje și dashboard."),
        ("REST API", "Stil de proiectare API care folosește resurse, URL-uri și metode HTTP. În VetCare, resursele principale sunt users, appointments și contact messages."),
        ("Endpoint", "Adresă concretă din API care execută o operație, de exemplu POST /api/auth/login sau GET /api/appointments/mine."),
        ("HTTP", "Protocolul prin care browserul și serverul schimbă cereri și răspunsuri. Metodele folosite în proiect sunt GET, POST, PUT și DELETE."),
        ("HTTP Client", "Componentă software care trimite cereri HTTP. În Angular, HttpClient este folosit pentru comunicarea dintre componente/servicii și backend."),
        ("Request", "Cererea trimisă de client către server. Poate conține metodă, URL, headere, token JWT și body JSON."),
        ("Response", "Răspunsul serverului către client. Include cod de status, headere și date JSON sau mesaj de eroare."),
        ("JSON", "Format text pentru schimb de date. VetCare îl folosește pentru request body și response body între Angular și Express."),
        ("CORS", "Mecanism de securitate al browserului care controlează accesul între origini diferite. Backend-ul permite frontend-ul local prin configurarea CORS."),
        ("Middleware", "Funcție Express executată înaintea rutelor finale. În VetCare, middleware-ul verifică JWT-ul și rolul de administrator."),
        ("Node.js", "Mediu de execuție JavaScript pe server. Rulează aplicația backend VetCare."),
        ("Express", "Framework web pentru Node.js, folosit la definirea rutelor API, middleware-urilor și handlerelor de eroare."),
        ("Controller", "Componentă backend care conține logica unei operații, cum ar fi register, login sau obținerea statisticilor."),
        ("Model", "Structură care descrie forma datelor aplicației. În VetCare, user.model.js creează obiectul utilizator cu UUID și câmpuri standard."),
        ("URL / Route", "Adresă sau rută prin care utilizatorul sau frontend-ul accesează o pagină ori un endpoint API."),
        ("Bază de date", "Sistem de stocare persistentă. VetCare folosește MySQL cu tabele pentru utilizatori, programări și mesaje."),
    ]
    add_table(doc, ["Termen", "Definiție și rol în proiect"], definition_rows, [3.2, 12.0], 8)
    add_heading(doc, "3.2 Frontend Angular", 2)
    for p in [
        "Frontend-ul este responsabil de experiența utilizatorului. În aplicația VetCare, Angular gestionează navigarea între pagini, afișarea calendarului de programări, formularele de login și register, profilul utilizatorului și dashboard-ul administrativ. O parte importantă a interfeței este feedback-ul imediat: butoane dezactivate în timpul trimiterii datelor, mesaje de eroare și evidențierea datei selectate.",
        "Componentele standalone reduc dependența de module globale și permit declararea explicită a importurilor necesare fiecărei pagini. Această abordare face structura proiectului mai ușor de urmărit pentru un proiect de licență, deoarece fiecare component are fișier TypeScript, HTML și CSS propriu.",
        "Serviciul de autentificare gestionează tokenul, utilizatorul curent și operațiile de login, register și logout. Separarea acestei logici într-un serviciu previne duplicarea codului în componente și permite extinderea ulterioară cu refresh token, stocare mai sigură sau integrare cu un provider extern.",
    ]:
        add_para(doc, p)
    add_heading(doc, "3.3 Backend Node.js și Express", 2)
    for p in [
        "Backend-ul expune un API REST sub prefixul /api. Fiecare grup de funcționalități are propriul fișier de rutare: auth.routes.js, appointments.routes.js, dashboard.routes.js și contact.routes.js. Această structură ajută la separarea responsabilităților și la localizarea rapidă a unei funcții.",
        "Express permite atașarea middleware-urilor globale, cum ar fi CORS, parsarea JSON și rate limiting. Înainte de rutele protejate, backend-ul verifică tokenul JWT și atașează informațiile despre utilizator la obiectul request. Pentru rutele de administrator, se verifică suplimentar rolul admin.",
        "Serverul tratează și cazurile de eroare. Dacă o rută nu există, răspunsul este 404 cu mesaj în limba română. Dacă apare o eroare internă, middleware-ul final returnează un mesaj generic, evitând expunerea detaliilor tehnice către client.",
    ]:
        add_para(doc, p)
    add_heading(doc, "3.4 Bază de date MySQL", 2)
    for p in [
        "Baza de date conține trei tabele: users, appointments și contact_messages. Tabela users păstrează conturile, tabela appointments păstrează programările, iar tabela contact_messages păstrează mesajele primite prin formularul public. Relația principală este users - appointments, de tip unu-la-mulți.",
        "Folosirea cheilor externe permite păstrarea integrității datelor. Dacă un utilizator este șters, programările sale sunt șterse automat prin ON DELETE CASCADE. Această regulă simplifică administrarea și previne apariția programărilor fără proprietar.",
    ]:
        add_para(doc, p)


def architecture(doc):
    add_heading(doc, "4. Arhitectura și proiectarea sistemului", 1)
    for p in [
        "Arhitectura aplicației VetCare este de tip client-server. Clientul este aplicația Angular rulată în browser, iar serverul este aplicația Node.js/Express care comunică cu baza de date MySQL. Comunicarea dintre frontend și backend se face prin cereri HTTP, iar datele sunt transmise în format JSON.",
        "Această arhitectură este potrivită pentru proiect deoarece separă clar prezentarea de logica de business. Frontend-ul se ocupă de interacțiunea cu utilizatorul, validări de bază și afișarea datelor. Backend-ul se ocupă de autentificare, validare finală, reguli de programare, acces la baza de date și securitate.",
    ]:
        add_para(doc, p)
    add_table(doc, ["Strat", "Responsabilități", "Exemple în VetCare"], [
        ("Prezentare", "Afișare UI, formulare, navigare", "Componente Angular: Home, Login, Programare, Dashboard"),
        ("Servicii client", "Comunicare HTTP și stare locală", "AuthService, HttpClient, token JWT"),
        ("API backend", "Rute, validări, reguli de business", "Express routes, middleware auth/admin"),
        ("Persistență", "Stocare și interogare date", "MySQL: users, appointments, contact_messages"),
    ], [3.2, 5.8, 6.2], 9)
    add_heading(doc, "4.1 Arhitectura generală", 2)
    for p in [
        "Figura 1 poate fi reprezentată conceptual printr-un flux în patru niveluri: utilizatorul interacționează cu browserul, browserul rulează aplicația Angular, Angular trimite cereri către API-ul Express, iar API-ul citește sau modifică datele în MySQL. Pentru operațiile protejate, cererea conține headerul Authorization cu tokenul JWT.",
        "În cazul autentificării, utilizatorul trimite email și parolă către backend. Serverul caută utilizatorul după email, compară parola cu hash-ul bcrypt și, dacă datele sunt corecte, emite un token semnat. Tokenul este păstrat pe client și transmis la cererile ulterioare. Middleware-ul verifică semnătura și extrage rolul.",
        "În cazul programării, interfața afișează calendarul, iar pentru o dată selectată cere backend-ului lista orelor ocupate. La trimiterea programării, serverul verifică data, ziua săptămânii, ora, serviciul, tipul animalului și conflictul de slot. Doar după aceste validări se inserează în baza de date.",
    ]:
        add_para(doc, p)
    add_heading(doc, "4.2 Organizarea frontend-ului", 2)
    component_rows = [
        ("home", "Pagina principală, prezentare servicii și motive de alegere"),
        ("about", "Prezentarea cabinetului și echipei medicale"),
        ("contact", "Formular public pentru mesaje"),
        ("login", "Autentificare și înregistrare"),
        ("programare", "Calendar și formular de programare"),
        ("profile", "Date utilizator și istoric programări"),
        ("dashboard", "Panou administrativ"),
        ("header/footer", "Navigare și informații generale"),
    ]
    add_table(doc, ["Componentă", "Rol"], component_rows, [4.2, 11.0], 9)
    for p in [
        "Structura componentelor este adecvată pentru o aplicație de dimensiune medie. Fiecare pagină are propriul fișier HTML pentru template, CSS pentru stilizare și TypeScript pentru logică. Această separare este utilă în mentenanță, deoarece modificarea designului nu necesită atingerea logicii de business.",
        "Rutarea este definită în app.routes.ts. Ruta pentru programare este protejată cu authGuard, iar ruta pentru dashboard este protejată atât cu authGuard, cât și cu adminGuard. Astfel, accesul este controlat încă din interfață, iar backend-ul aplică verificarea finală.",
    ]:
        add_para(doc, p)
    add_heading(doc, "4.3 Organizarea backend-ului", 2)
    backend_rows = [
        ("server.js", "Configurare Express, middleware-uri globale, rute și pornire server"),
        ("routes", "Definirea endpointurilor pe domenii funcționale"),
        ("controllers", "Logică pentru autentificare și dashboard"),
        ("middleware", "Verificare JWT și rol admin"),
        ("config/db.js", "Adaptor pentru interacțiunea cu MySQL"),
        ("utils/validation.js", "Funcții de validare și sanitizare"),
    ]
    add_table(doc, ["Fișier/folder", "Responsabilitate"], backend_rows, [4.2, 11.0], 9)
    for p in [
        "Organizarea backend-ului urmărește separarea regulilor de business de configurarea serverului. Rutele primesc cererile HTTP și apelează controllere sau logică specifică. Middleware-ul de autentificare este reutilizat pentru toate rutele protejate.",
        "Un aspect important este faptul că userId-ul programării nu este primit din corpul cererii, ci este extras din token. Această decizie previne impersonarea, adică situația în care un utilizator ar încerca să creeze o programare în numele altui utilizator modificând manual corpul cererii.",
    ]:
        add_para(doc, p)


def database_api(doc):
    add_heading(doc, "5. Proiectarea bazei de date și a API-ului", 1)
    for p in [
        "Proiectarea bazei de date a urmărit păstrarea unui model simplu, ușor de explicat și suficient pentru fluxurile implementate. Entitățile principale sunt utilizatorul, programarea și mesajul de contact. Deși într-un sistem complet ar exista și entități pentru animale, fișe medicale, tratamente și facturi, acestea au fost lăsate ca direcții de extindere.",
        "Cheile primare sunt de tip UUID stocat ca VARCHAR(36). Alegerea UUID-urilor oferă identificatori care nu sunt ușor de ghicit și care pot fi generați independent. Pentru dimensiunea aplicației, costul suplimentar de stocare nu reprezintă o problemă.",
    ]:
        add_para(doc, p)
    add_heading(doc, "5.1 Schema relațională", 2)
    db_rows = [
        ("users", "id, email, password, name, role, createdAt", "Conturi și roluri"),
        ("appointments", "id, userId, date, service, animalType, message, status, createdAt", "Programări"),
        ("contact_messages", "id, firstName, lastName, email, phone, animalType, subject, message, read, createdAt", "Mesaje contact"),
    ]
    add_table(doc, ["Tabel", "Câmpuri principale", "Rol"], db_rows, [3.2, 8.0, 4.0], 8)
    for p in [
        "Tabela users conține emailul unic, hash-ul parolei și rolul. Rolul poate fi user sau admin. Această informație este introdusă în tokenul JWT și este folosită ulterior pentru accesul la rutele administrative.",
        "Tabela appointments include data și ora programării, serviciul, tipul animalului și statusul. Statusul poate fi pending, confirmed sau cancelled. Prin păstrarea statusului, aplicația nu este obligată să șteargă programările anulate, ceea ce permite păstrarea istoricului.",
        "Tabela contact_messages nu are cheie externă către users deoarece mesajele pot fi trimise și de vizitatori care nu au cont. Această decizie reflectă realitatea unui formular public, unde emailul introdus este doar o informație de contact și nu neapărat un cont existent.",
    ]:
        add_para(doc, p)
    add_heading(doc, "5.2 Tipuri de chei folosite în scriptul MySQL", 2)
    key_rows = [
        ("Primary Key", "Cheie primară", "Identifică unic fiecare rând dintr-un tabel. În VetCare, câmpul id este PRIMARY KEY în users, appointments și contact_messages."),
        ("Foreign Key", "Cheie externă", "Leagă două tabele și impune integritatea referențială. appointments.userId referă users.id."),
        ("Unique Key", "Cheie unică", "Nu permite valori duplicate într-o coloană. În tabela users, email este UNIQUE pentru a preveni conturi duplicate."),
        ("Index / Key", "Index de căutare", "Accelerează interogările pe coloane folosite frecvent. În proiect apar indecși pe appointments.date, appointments.userId, contact_messages.email și contact_messages.read."),
        ("Surrogate Key", "Cheie surogat", "Cheie artificială, fără semnificație de business. UUID-ul id este o cheie surogat, deoarece nu descrie utilizatorul sau programarea."),
        ("Natural Key", "Cheie naturală", "Câmp cu semnificație reală care poate identifica o entitate. Emailul poate fi privit ca o cheie naturală, dar în proiect nu este primary key, ci unique key."),
        ("Composite Key", "Cheie compusă", "Cheie formată din mai multe coloane. Scriptul actual nu folosește chei compuse, deoarece fiecare tabel are id UUID separat."),
        ("ON DELETE CASCADE", "Acțiune referențială", "Dacă un utilizator este șters, programările sale sunt șterse automat. Regula este definită pe cheia externă appointments.userId."),
    ]
    add_table(doc, ["Termen SQL", "Denumire", "Explicație în VetCare"], key_rows, [3.0, 3.4, 8.8], 8)
    for p in [
        "În scriptul MySQL, cuvântul KEY este folosit în mai multe contexte. PRIMARY KEY definește cheia principală a tabelului, UNIQUE KEY definește o constrângere de unicitate, iar KEY simplu definește un index folosit pentru performanță. De exemplu, KEY date_idx (date) nu impune unicitate, ci ajută la căutarea rapidă a programărilor după dată.",
        "Cheia externă appointments_ibfk_1 este importantă deoarece exprimă relația dintre utilizatori și programări. O programare trebuie să aparțină unui utilizator existent. Această regulă împiedică introducerea unor date orfane și păstrează coerența bazei de date.",
        "Alegerea UUID-ului ca primary key este o decizie de proiectare. Alternativa ar fi fost un id numeric auto-increment. UUID-ul este mai greu de ghicit și este potrivit pentru aplicații care pot fi extinse ulterior, deși ocupă mai mult spațiu decât un număr întreg.",
    ]:
        add_para(doc, p)
    add_code_block(doc, """
PRIMARY KEY (`id`)
UNIQUE KEY `email` (`email`)
KEY `userId` (`userId`)
KEY `date_idx` (`date`)
CONSTRAINT `appointments_ibfk_1`
  FOREIGN KEY (`userId`) REFERENCES `users` (`id`) ON DELETE CASCADE
""")
    add_heading(doc, "5.3 API REST", 2)
    endpoints = [
        ("POST", "/api/auth/register", "Public", "Creează un cont nou"),
        ("POST", "/api/auth/login", "Public", "Autentifică utilizatorul și returnează token"),
        ("GET", "/api/auth/me", "User", "Returnează utilizatorul curent"),
        ("PUT", "/api/auth/profile", "User", "Actualizează numele"),
        ("PUT", "/api/auth/change-password", "User", "Schimbă parola"),
        ("GET", "/api/appointments/occupied", "User", "Returnează ore ocupate pentru o dată"),
        ("GET", "/api/appointments/mine", "User", "Returnează programările proprii"),
        ("POST", "/api/appointments", "User", "Creează programare"),
        ("DELETE", "/api/appointments/:id", "User/Admin", "Anulează programare"),
        ("POST", "/api/contact", "Public", "Trimite mesaj de contact"),
        ("GET", "/api/dashboard/stats", "Admin", "Statistici globale"),
        ("GET", "/api/dashboard/users", "Admin", "Lista utilizatorilor"),
        ("PUT", "/api/dashboard/users/:id/role", "Admin", "Schimbă rol"),
        ("DELETE", "/api/dashboard/users/:id", "Admin", "Șterge utilizator"),
        ("GET", "/api/dashboard/appointments", "Admin", "Lista tuturor programărilor"),
        ("PUT", "/api/dashboard/appointments/:id/status", "Admin", "Schimbă status programare"),
        ("GET", "/api/dashboard/messages", "Admin", "Lista mesajelor"),
    ]
    add_table(doc, ["Metodă", "Endpoint", "Acces", "Descriere"], endpoints, [2.0, 5.3, 2.5, 5.4], 8)
    for p in [
        "API-ul respectă un stil REST pragmatic. Resursele sunt grupate după domenii funcționale, iar metodele HTTP sunt folosite în mod intuitiv: POST pentru creare, GET pentru citire, PUT pentru actualizare și DELETE pentru anulare sau ștergere. Răspunsurile de eroare au un format simplu, cu proprietatea message.",
        "Separarea endpointurilor publice de cele protejate este esențială. Formularul de contact și autentificarea sunt publice, dar programările și dashboard-ul necesită token. Dashboard-ul necesită și rol admin. Această structură previne accesarea accidentală sau intenționată a datelor administrative de către utilizatori obișnuiți.",
    ]:
        add_para(doc, p)
    add_heading(doc, "5.4 Validări și reguli de business", 2)
    for p in [
        "Regulile de business sunt implementate pe server, chiar dacă unele sunt reflectate și în interfață. De exemplu, calendarul poate marca weekendurile ca indisponibile, dar backend-ul verifică din nou ziua săptămânii. Această dublă verificare este importantă deoarece datele trimise de client nu pot fi considerate de încredere.",
        "La crearea unei programări, serverul verifică existența câmpurilor obligatorii, apartenența serviciului și a tipului de animal la listele acceptate, validitatea datei, faptul că data nu este în trecut, intervalul orar și conflictul de slot. Doar după aceste validări se creează înregistrarea.",
        "Pentru autentificare, backend-ul normalizează emailul la litere mici și verifică formatul. Parola trebuie să aibă cel puțin opt caractere și să conțină literă și cifră. Aceste reguli nu înlocuiesc politici complexe de securitate, dar sunt suficiente pentru un proiect de licență și demonstrează atenție la validarea inputului.",
    ]:
        add_para(doc, p)


def implementation(doc):
    add_heading(doc, "6. Implementarea aplicației VetCare", 1)
    for p in [
        "Implementarea aplicației a fost realizată incremental, pornind de la schema bazei de date și de la fluxurile principale. Backend-ul a fost creat pentru a expune endpointurile necesare, iar frontend-ul a fost dezvoltat pentru a consuma aceste endpointuri prin HttpClient. În timpul implementării, s-a urmărit păstrarea mesajelor în limba română și coerența dintre validarea din interfață și validarea server-side.",
        "Proiectul este împărțit în două directoare principale: vetcare-backend și vetcare-frontend. Această separare permite rularea independentă a serverului API și a aplicației Angular. Backend-ul rulează implicit pe portul 3000, iar frontend-ul pe portul 4200.",
    ]:
        add_para(doc, p)
    add_heading(doc, "6.1 Implementarea autentificării", 2)
    for p in [
        "La înregistrare, parola este hash-uită cu bcrypt înainte de a fi stocată. În răspunsul către client, parola este eliminată din obiectul utilizatorului prin funcția safeUser. La login, parola primită este comparată cu hash-ul din baza de date, iar dacă verificarea reușește, serverul semnează un token JWT.",
        "Tokenul conține id-ul utilizatorului, emailul și rolul. Aceste informații sunt suficiente pentru verificarea accesului la rutele protejate. Expirarea tokenului este configurabilă prin variabila JWT_EXPIRES_IN, cu valoare implicită de șapte zile.",
    ]:
        add_para(doc, p)
    add_code_block(doc, """
function signToken(user) {
  return jwt.sign(
    { id: user.id, email: user.email, role: user.role },
    process.env.JWT_SECRET,
    { expiresIn: process.env.JWT_EXPIRES_IN || '7d' }
  );
}
""")
    add_heading(doc, "6.2 Implementarea programărilor", 2)
    for p in [
        "Componenta Programare gestionează calendarul, data selectată, ora selectată și formularul de programare. Calendarul generează zilele lunii curente și marchează zilele trecute, weekendurile și ziua curentă. La selectarea unei zile valide, componenta cere backend-ului sloturile ocupate.",
        "Lista orelor este predefinită din 30 în 30 de minute între 08:00 și 19:30. Butoanele pentru ore ocupate sunt dezactivate, ceea ce reduce riscul alegerii unui interval indisponibil. Totuși, serverul verifică din nou conflictul pentru a evita situația în care două cereri ajung aproape simultan.",
        "La trimiterea formularului, componenta construiește un obiect Date care combină ziua selectată și ora selectată. Backend-ul primește data în format ISO, o validează și o transformă în format compatibil MySQL.",
    ]:
        add_para(doc, p)
    add_code_block(doc, """
const conflict = all.find(a => {
  if (!a.date || a.status === 'cancelled') return false;
  return new Date(a.date).getTime() === apptDate.getTime();
});
if (conflict) {
  return res.status(409).json({ message: 'Slotul este deja ocupat.' });
}
""")
    add_heading(doc, "6.3 Implementarea dashboard-ului administrativ", 2)
    for p in [
        "Dashboard-ul este disponibil doar administratorilor. Acesta oferă statistici globale, lista utilizatorilor, lista programărilor și lista mesajelor de contact. Prin intermediul dashboard-ului, administratorul poate modifica roluri, poate confirma sau anula programări și poate marca mesajele ca citite.",
        "Backend-ul protejează toate rutele dashboard prin requireAdmin. Această funcție apelează mai întâi requireAuth și apoi verifică rolul din payload-ul tokenului. Dacă rolul nu este admin, răspunsul este 403.",
        "Statisticile sunt utile pentru o vedere rapidă asupra activității cabinetului: numărul total de utilizatori, numărul de administratori, programări pending, confirmed sau cancelled și mesajele necitite. Aceste date pot fi extinse ulterior cu grafice, filtre pe intervale de timp sau rapoarte exportabile.",
    ]:
        add_para(doc, p)
    add_heading(doc, "6.4 Interfața cu utilizatorul", 2)
    for p in [
        "Interfața aplicației urmărește un stil prietenos, potrivit unui cabinet veterinar. Paginile publice au rol de prezentare și includ informații despre servicii, echipă și contact. Paginile protejate au un caracter mai operațional, în special programarea și profilul.",
        "Un principiu important al interfeței este reducerea ambiguității. Butoanele de acțiune folosesc texte clare, mesajele de eroare sunt în română, iar stările de încărcare sunt indicate prin textul „Se trimite...”. În calendar, zilele indisponibile sunt marcate vizual, iar ora selectată este evidențiată.",
        "Imaginile aplicației sunt stocate în src/assets/images și sunt referite prin path-uri relative, astfel încât Angular să le copieze în build. Pentru versiunea curentă au fost adăugate imagini fără filigran, descărcate din surse gratuite, cu fișier separat pentru surse.",
    ]:
        add_para(doc, p)
    add_heading(doc, "6.5 Securitate aplicată", 2)
    for p in [
        "Securitatea aplicației este abordată pe mai multe niveluri. La nivel de autentificare, parolele nu sunt stocate în clar, iar tokenurile sunt semnate. La nivel de acces, middleware-ul verifică tokenul și rolul. La nivel de input, datele sunt validate și sanitizate. La nivel de trafic local, CORS este configurat pentru frontend-ul așteptat.",
        "Rate limiting-ul este aplicat rutelor de login și register, cu limita de zece încercări în cincisprezece minute. Această măsură nu oprește toate atacurile, dar reduce riscul încercărilor brute-force simple și demonstrează o preocupare pentru protejarea endpointurilor sensibile.",
        "Pentru prevenirea SQL injection, proiectul folosește mysql2/promise și interogări parametrizate în adaptorul de bază de date. De asemenea, userId-ul nu este acceptat din corpul cererii de programare, ci este extras din token, ceea ce reduce riscul accesului neautorizat la datele altui utilizator.",
    ]:
        add_para(doc, p)


def testing(doc):
    add_heading(doc, "7. Testare, validare și evaluare", 1)
    for p in [
        "Testarea unei aplicații web full-stack trebuie să acopere atât interfața, cât și API-ul și baza de date. Pentru VetCare, testarea s-a concentrat pe fluxurile critice: înregistrare, autentificare, programare, anulare, acces dashboard și trimitere mesaj de contact. În plus, au fost verificate cazurile negative, cum ar fi parole slabe, date invalide sau acces neautorizat.",
        "Validarea este importantă deoarece aplicația primește date de la utilizatori. Chiar dacă interfața limitează alegerile posibile, backend-ul trebuie să verifice independent fiecare cerere. Această abordare este obligatorie în aplicațiile web, unde clientul poate fi modificat sau ocolit.",
    ]:
        add_para(doc, p)
    rows = [
        ("T01", "Register cu date valide", "Email nou, parolă puternică", "Cont creat, token returnat"),
        ("T02", "Register email duplicat", "Email existent", "Răspuns 409"),
        ("T03", "Login valid", "Email și parolă corecte", "Token și user fără parolă"),
        ("T04", "Login invalid", "Parolă greșită", "Răspuns 401"),
        ("T05", "Programare zi lucrătoare", "Dată viitoare, slot liber", "Programare creată"),
        ("T06", "Programare weekend", "Sâmbătă sau duminică", "Răspuns 400"),
        ("T07", "Programare în trecut", "Dată anterioară", "Răspuns 400"),
        ("T08", "Slot ocupat", "Aceeași dată și oră", "Răspuns 409"),
        ("T09", "Acces profil fără token", "Header lipsă", "Răspuns 401"),
        ("T10", "Acces dashboard ca user", "Token user", "Răspuns 403"),
        ("T11", "Trimite mesaj contact", "Date publice valide", "Mesaj stocat"),
        ("T12", "Marcare mesaj citit", "Token admin", "Status read actualizat"),
    ]
    add_table(doc, ["ID", "Scenariu", "Date de intrare", "Rezultat așteptat"], rows, [1.5, 4.6, 4.4, 4.7], 8)
    add_heading(doc, "7.1 Testarea funcțională", 2)
    for p in [
        "Testarea funcțională urmărește dacă aplicația face ceea ce promite cerințele. În cazul VetCare, un flux complet începe cu înregistrarea unui utilizator, continuă cu autentificarea, selectarea unei programări și vizualizarea acesteia în profil. Administratorul poate apoi vedea programarea în dashboard și îi poate schimba statusul.",
        "Un scenariu important este verificarea sloturilor ocupate. După crearea unei programări, endpointul /api/appointments/occupied trebuie să returneze ora respectivă pentru data aleasă. Interfața trebuie să dezactiveze butonul corespunzător, iar backend-ul trebuie să refuze o nouă programare pe același slot.",
        "Testarea mesajelor de contact verifică faptul că formularul public funcționează fără autentificare, dar gestionarea mesajelor este disponibilă doar administratorului. Această separare confirmă că aplicația combină corect accesul public cu administrarea protejată.",
    ]:
        add_para(doc, p)
    add_heading(doc, "7.2 Testarea securității", 2)
    for p in [
        "Testarea securității a urmărit verificarea rutelor protejate și a validărilor. Cererile fără header Authorization către rutele /api/auth/me, /api/appointments/mine sau /api/dashboard/stats trebuie să primească răspuns 401. Cererile cu token de utilizator obișnuit către dashboard trebuie să primească 403.",
        "Pentru autentificare, parolele invalide nu trebuie să permită login, iar mesajul returnat nu trebuie să indice dacă emailul sau parola a fost greșită. Această alegere reduce informațiile disponibile pentru un atacator.",
        "Pentru parole, aplicația folosește bcrypt cu factor de lucru 10. Conform recomandărilor OWASP, bcrypt este o variantă acceptabilă pentru stocarea parolelor atunci când se folosește un factor de lucru suficient și parolele nu sunt stocate în clar.",
    ]:
        add_para(doc, p)
    add_heading(doc, "7.3 Evaluarea soluției", 2)
    for p in [
        "Soluția implementată îndeplinește obiectivele definite: oferă o interfață publică, autentificare, programări online, profil utilizator, formular de contact și dashboard administrativ. Din punct de vedere tehnic, aplicația demonstrează integrarea dintre frontend, backend și baza de date.",
        "Punctele forte ale aplicației sunt simplitatea arhitecturii, claritatea fluxurilor și existența unor măsuri de securitate de bază. Punctele care pot fi îmbunătățite sunt lipsa unei entități separate pentru animale, lipsa notificărilor prin email, lipsa unui calendar administrativ avansat și lipsa testelor automate extinse.",
        "Într-un context real, aplicația ar trebui completată cu politici de backup, logare structurată, HTTPS, protecție CSRF în funcție de metoda de stocare a tokenului, monitorizare și management al erorilor. Pentru scopul lucrării, însă, aplicația oferă un nucleu coerent și demonstrabil.",
    ]:
        add_para(doc, p)


def conclusions(doc):
    add_heading(doc, "8. Concluzii și direcții viitoare", 1)
    for p in [
        "Lucrarea a prezentat procesul de proiectare și implementare a aplicației VetCare, o platformă web pentru gestionarea unui cabinet veterinar. Proiectul demonstrează utilizarea unei arhitecturi full-stack moderne, cu Angular pe client, Node.js/Express pe server și MySQL pentru persistența datelor.",
        "Obiectivul principal a fost atins prin realizarea unui sistem funcțional care permite vizitatorilor să interacționeze cu cabinetul, utilizatorilor să se programeze online și administratorilor să gestioneze informațiile operaționale. Implementarea a evidențiat importanța separării responsabilităților, a validării server-side și a protejării rutelor sensibile.",
        "Din punct de vedere academic, proiectul valorifică noțiuni de programare web, baze de date, inginerie software, securitate și testare. Contribuția proprie este vizibilă în definirea fluxurilor, implementarea programării online, proiectarea modelului de date și documentarea soluției.",
        "Ca direcții viitoare, aplicația poate fi extinsă cu evidența animalelor, fișe medicale, încărcare documente, notificări automate prin email sau SMS, calendar avansat pentru administratori, rapoarte statistice, export PDF și integrarea plăților online. O altă direcție importantă este adăugarea testelor automate end-to-end și a unui pipeline de deploy.",
        "În concluzie, VetCare reprezintă o soluție educațională completă pentru un scenariu realist. Aplicația nu înlocuiește un sistem clinic profesional, dar oferă o bază solidă pentru digitalizarea operațiilor simple ale unui cabinet veterinar și pentru continuarea dezvoltării într-un produs mai complex.",
    ]:
        add_para(doc, p)


def bibliography(doc):
    add_heading(doc, "Bibliografie", 1)
    refs = [
        "[1] Facultatea de Matematică și Informatică, Universitatea „Ovidius” din Constanța, Ghid de elaborare și susținere a lucrării de licență și disertație, 2025, https://fmi.univ-ovidius.ro/wp-content/uploads/2025/05/Ghid-de-elaborare-a-lucrarii-de-finalizare-studii-FMI2025.pdf",
        "[2] Facultatea de Matematică și Informatică, Organizarea Examenului de Licență, https://fmi.univ-ovidius.ro/en/academic/studii-de-licenta/licenta-documente/",
        "[3] Angular Documentation, Define routes, https://angular.dev/guide/routing/define-routes",
        "[4] Angular Documentation, Component API, https://angular.dev/api/core/Component",
        "[5] Express.js Documentation, Routing, https://expressjs.com/en/guide/routing.html",
        "[6] Node.js Documentation, Introduction to Node.js, https://nodejs.org/en/learn/getting-started/introduction-to-nodejs",
        "[7] MySQL 8.4 Reference Manual, FOREIGN KEY Constraints, https://dev.mysql.com/doc/refman/8.4/en/create-table-foreign-keys.html",
        "[8] RFC Editor, RFC 7519: JSON Web Token (JWT), https://www.rfc-editor.org/rfc/rfc7519",
        "[9] OWASP Cheat Sheet Series, Password Storage Cheat Sheet, https://cheatsheetseries.owasp.org/cheatsheets/Password_Storage_Cheat_Sheet.html",
        "[10] MDN Web Docs, Cross-Origin Resource Sharing (CORS), https://developer.mozilla.org/en-US/docs/Web/HTTP/CORS",
        "[11] VetCare, README.md și documentația tehnică locală a proiectului: API_ENDPOINTS.md, ER_DIAGRAM.md, vetcare_schema.sql.",
    ]
    for ref in refs:
        add_para(doc, ref, align=WD_ALIGN_PARAGRAPH.LEFT)


def code_screenshots_appendix(doc):
    add_heading(doc, "Anexa F. Capturi de cod și scripturi relevante", 1)
    add_para(doc, "În această anexă sunt incluse capturi generate din fișierele sursă ale proiectului. Ele arată concret unde apar HTML, CSS, TypeScript, controllere, modele, rute URL și scriptul SQL al bazei de date.")
    code_dir = WORK_DIR / "code_screenshots"
    items = [
        (BASE_DIR / "vetcare_schema.sql", "sql_schema.png", "Script SQL - crearea bazei de date VetCare", "Figura F.1. Script SQL pentru baza de date"),
        (BASE_DIR / "vetcare-frontend/src/app/components/programare/programare.component.ts", "typescript_programare.png", "TypeScript - componenta Programare", "Figura F.2. Logică TypeScript pentru programări"),
        (BASE_DIR / "vetcare-frontend/src/app/components/programare/programare.component.html", "html_programare.png", "HTML - template Programare", "Figura F.3. Template HTML pentru calendarul de programări"),
        (BASE_DIR / "vetcare-frontend/src/app/components/programare/programare.component.css", "css_programare.png", "CSS - stilizare Programare", "Figura F.4. Stilizare CSS pentru pagina Programare"),
        (BASE_DIR / "vetcare-backend/src/controllers/auth.controller.js", "controller_auth.png", "Controller - autentificare", "Figura F.5. Controller backend pentru register/login"),
        (BASE_DIR / "vetcare-backend/src/models/user.model.js", "model_user.png", "Model - utilizator", "Figura F.6. Model backend pentru utilizator"),
        (BASE_DIR / "vetcare-frontend/src/app/app.routes.ts", "routes_frontend.png", "URL/Routing - rute Angular", "Figura F.7. Definirea rutelor frontend"),
        (BASE_DIR / "vetcare-backend/src/routes/appointments.routes.js", "routes_backend.png", "URL/API - rute programări", "Figura F.8. Endpointuri backend pentru programări"),
        (BASE_DIR / "vetcare-frontend/src/app/services/auth.service.ts", "httpclient_auth_service.png", "Angular HttpClient - serviciu autentificare", "Figura F.9. Folosirea HttpClient pentru apeluri API"),
        (BASE_DIR / "vetcare-backend/src/middleware/auth.middleware.js", "middleware_auth.png", "Middleware - verificare JWT", "Figura F.10. Middleware Express pentru autentificare și rol admin"),
        (BASE_DIR / "vetcare-backend/src/server.js", "server_api_config.png", "Server API - Express, CORS, rate limit", "Figura F.11. Configurarea API-ului Express"),
        (BASE_DIR / "vetcare-backend/src/config/db.js", "db_adapter.png", "Adaptor bază de date - MySQL", "Figura F.12. Funcții de acces la baza de date"),
    ]
    for source, out_name, title, caption in items:
        image = make_code_image(source, code_dir / out_name, title)
        add_figure(doc, image, caption, width_cm=15.4)


def appendices(doc):
    add_heading(doc, "Anexa A. Instrucțiuni de instalare și rulare", 1)
    for p in [
        "Pentru rularea aplicației sunt necesare Node.js, npm și MySQL. Baza de date se creează folosind scriptul vetcare_schema.sql. Backend-ul se configurează prin fișierul .env, unde se setează portul, secretul JWT, datele de conectare la MySQL și URL-ul frontend-ului.",
        "Pașii generali sunt: crearea bazei de date, instalarea dependențelor backend, popularea opțională cu date demo, pornirea backend-ului, instalarea dependențelor frontend și pornirea aplicației Angular. Pentru testare rapidă, proiectul include conturi demo generate de seed.js.",
    ]:
        add_para(doc, p)
    add_code_block(doc, """
mysql -u root -p < vetcare_schema.sql
cd vetcare-backend
npm install
npm run seed
npm start

cd ../vetcare-frontend
npm install
npm start
""")
    add_heading(doc, "Anexa B. Scenarii de test detaliate", 1)
    detailed = []
    for i in range(1, 31):
        area = ["Autentificare", "Programări", "Dashboard", "Contact", "Profil"][i % 5]
        action = [
            "verificarea răspunsului pentru date valide",
            "verificarea validării pentru câmpuri lipsă",
            "verificarea accesului fără token",
            "verificarea accesului cu rol insuficient",
            "verificarea persistării datelor în MySQL",
        ][i % 5]
        expected = [
            "răspuns 200 sau 201 și payload JSON corect",
            "răspuns 400 cu mesaj în limba română",
            "răspuns 401 și blocarea operației",
            "răspuns 403 pentru resurse administrative",
            "înregistrarea este creată, actualizată sau filtrată corect",
        ][i % 5]
        detailed.append((f"TB{i:02d}", area, action, expected))
    add_table(doc, ["ID", "Zonă", "Acțiune", "Rezultat așteptat"], detailed, [1.6, 3.0, 5.4, 5.2], 8)
    add_heading(doc, "Anexa C. Fragment de schemă SQL", 1)
    add_code_block(doc, """
CREATE TABLE `appointments` (
  `id`         VARCHAR(36)  NOT NULL,
  `userId`     VARCHAR(36)           DEFAULT NULL,
  `date`       DATETIME              DEFAULT NULL,
  `service`    VARCHAR(255)          DEFAULT NULL,
  `animalType` VARCHAR(255)          DEFAULT '',
  `message`    VARCHAR(1000)         DEFAULT '',
  `status`     ENUM('pending','confirmed','cancelled') DEFAULT 'pending',
  `createdAt`  DATETIME              DEFAULT CURRENT_TIMESTAMP,
  PRIMARY KEY (`id`),
  KEY `userId` (`userId`),
  KEY `date_idx` (`date`),
  CONSTRAINT `appointments_ibfk_1`
    FOREIGN KEY (`userId`) REFERENCES `users` (`id`) ON DELETE CASCADE
);
""")
    add_heading(doc, "Anexa D. Întrebări posibile pentru susținere", 1)
    questions = [
        "Care este rolul JWT în aplicație și ce informații sunt incluse în token?",
        "De ce validările importante sunt realizate pe backend, chiar dacă există validări în frontend?",
        "Cum este prevenită programarea pe un slot deja ocupat?",
        "Care este diferența dintre rolul user și rolul admin în VetCare?",
        "De ce parolele nu sunt stocate în clar în baza de date?",
        "Ce avantaje oferă separarea frontend/backend?",
        "Cum poate fi extinsă aplicația pentru fișe medicale ale animalelor?",
        "Ce rol au cheile externe în schema MySQL?",
        "Ce limitări are aplicația în forma curentă?",
        "Ce teste automate ar fi utile pentru o versiune de producție?",
    ]
    for q in questions:
        add_bullet(doc, q)
    add_para(doc, "Aceste întrebări sunt orientative și urmăresc pregătirea pentru discuția cu comisia. Răspunsurile trebuie susținute prin demonstrația aplicației și prin referire la capitolele tehnice ale lucrării.")


def expand_content(doc):
    """Add natural student-style observations without changing the formal structure."""
    add_heading(doc, "Anexa E. Observații personale privind proiectarea și implementarea", 1)
    add_para(doc, "Această anexă completează partea tehnică a lucrării cu observații scrise într-un stil mai apropiat de modul în care am lucrat efectiv la aplicație. Am inclus aici explicații despre alegeri, dificultăți și posibile îmbunătățiri, deoarece la susținere aceste detalii sunt utile pentru a arăta că proiectul a fost înțeles, nu doar implementat.")

    sections = [
        (
            "De ce am separat frontend-ul de backend",
            [
                "Am ales să separ aplicația în două părți, frontend și backend, deoarece așa am putut urmări mai clar ce se întâmplă în fiecare zonă. În frontend m-am ocupat de pagini, formulare, calendar și afișarea mesajelor. În backend am păstrat autentificarea, verificarea tokenului, validarea datelor și accesul la baza de date.",
                "Această separare m-a ajutat și la depanare. Dacă o pagină nu afișa datele corect, puteam verifica separat dacă problema era în componenta Angular sau în răspunsul API-ului. Pentru un proiect de licență, această claritate este importantă, pentru că fiecare parte poate fi explicată separat în fața comisiei.",
                "Un alt motiv este extinderea ulterioară. Dacă aplicația ar fi transformată într-un produs real, backend-ul ar putea fi folosit și de o aplicație mobilă, nu doar de interfața Angular. În acest caz, API-ul ar rămâne nucleul comun al sistemului."
            ],
        ),
        (
            "Cum am gândit autentificarea",
            [
                "Autentificarea a fost una dintre primele funcționalități importante. Am vrut ca utilizatorul să poată crea cont, să se autentifice și apoi să aibă acces la pagini protejate. Pentru acest lucru am folosit JWT, deoarece tokenul poate fi trimis ușor în headerul Authorization la fiecare cerere către server.",
                "În token am păstrat doar informațiile necesare: id-ul utilizatorului, emailul și rolul. Nu am pus parola sau alte date sensibile. Parola este verificată doar la login, iar în baza de date este stocată sub formă de hash bcrypt. Astfel, chiar dacă cineva ar vedea tabela users, nu ar vedea parolele în clar.",
                "Pentru roluri am păstrat o variantă simplă: user și admin. În aplicația actuală este suficient, pentru că userul face programări și vede profilul propriu, iar adminul gestionează utilizatori, programări și mesaje. Dacă proiectul ar fi extins, s-ar putea adăuga roluri precum medic veterinar sau recepționer."
            ],
        ),
        (
            "Probleme întâlnite la programări",
            [
                "Partea de programări a fost mai interesantă decât un formular simplu, deoarece nu era suficient să salvez o dată în baza de date. A trebuit să verific dacă data este validă, dacă nu este în trecut, dacă nu este weekend și dacă ora este în intervalul de lucru.",
                "O problemă importantă a fost slotul ocupat. Chiar dacă în interfață pot dezactiva o oră ocupată, verificarea reală trebuie făcută pe server. Utilizatorul poate modifica cererea sau poate trimite direct un request către API. De aceea, backend-ul caută în programările existente și returnează eroare dacă există conflict.",
                "Am ales intervale din 30 în 30 de minute pentru că sunt ușor de afișat în calendar și potrivite pentru o aplicație demonstrativă. Într-un cabinet real, durata unei consultații poate depinde de serviciu. De exemplu, o consultație simplă poate dura 30 de minute, iar o intervenție chirurgicală mult mai mult."
            ],
        ),
        (
            "De ce am folosit MySQL și chei UUID",
            [
                "Pentru baza de date am folosit MySQL deoarece datele proiectului sunt relaționale. Un utilizator poate avea mai multe programări, iar mesajele de contact trebuie păstrate separat. Modelul cu tabele este ușor de explicat și de verificat prin script SQL.",
                "Am folosit UUID pentru câmpurile id. O variantă mai simplă ar fi fost id numeric auto-increment, dar UUID-ul are avantajul că nu poate fi ghicit ușor. De exemplu, un utilizator nu poate presupune că următoarea programare are id-ul 15 sau 16. Pentru dimensiunea acestui proiect, spațiul ocupat în plus nu este o problemă.",
                "Cheia externă dintre appointments și users este importantă pentru integritatea datelor. O programare trebuie să aparțină unui utilizator existent. Regula ON DELETE CASCADE face ca programările unui utilizator să fie eliminate automat dacă acel utilizator este șters."
            ],
        ),
        (
            "Ce am urmărit în interfață",
            [
                "La interfață am încercat să păstrez un stil clar și prietenos, potrivit pentru un cabinet veterinar. Nu am vrut ca aplicația să arate ca un formular tehnic rece, ci ca un site pe care un client îl poate folosi fără multe explicații.",
                "Am folosit pagini separate pentru Home, About, Contact, Login, Programare, Profile și Dashboard. Această împărțire face aplicația mai ușor de urmărit. De exemplu, componenta Programare conține calendarul și formularul, iar Dashboard este separat pentru administrator.",
                "Mesajele sunt în limba română pentru că aplicația este gândită pentru un cabinet local. Acest detaliu contează, deoarece erorile de tip «Completează toate câmpurile» sau «Slotul este deja ocupat» sunt mai utile pentru utilizator decât mesaje tehnice în engleză."
            ],
        ),
        (
            "Ce aș îmbunătăți într-o versiune următoare",
            [
                "Prima îmbunătățire ar fi adăugarea unei tabele animals. În versiunea actuală se salvează doar tipul animalului, cum ar fi câine sau pisică. Într-o aplicație reală, fiecare utilizator ar putea avea mai multe animale, cu nume, vârstă, rasă și istoric medical.",
                "A doua îmbunătățire ar fi notificările. După confirmarea unei programări, utilizatorul ar putea primi email sau SMS. De asemenea, administratorul ar putea primi o notificare când apare un mesaj nou de contact.",
                "A treia direcție ar fi testarea automată mai serioasă. În proiect există scenarii de test descrise, dar într-o versiune de producție aș adăuga teste pentru API, teste pentru componentele Angular și teste end-to-end pentru fluxul de programare."
            ],
        ),
        (
            "Ce pot explica la susținere",
            [
                "La susținere pot explica traseul unei cereri complete. De exemplu, când utilizatorul face login, Angular trimite email și parolă către endpointul /api/auth/login. Backend-ul verifică datele, compară parola cu hash-ul bcrypt și returnează tokenul JWT. Apoi frontend-ul folosește tokenul pentru cererile protejate.",
                "Pot explica și diferența dintre validarea din frontend și validarea din backend. Frontend-ul ajută utilizatorul să nu greșească, dar backend-ul este cel care decide dacă datele sunt acceptate. Din acest motiv, regulile importante sunt duplicate sau confirmate pe server.",
                "Un alt punct bun de explicat este schema bazei de date. Tabela users este legată de appointments prin userId, iar contact_messages este separată deoarece mesajele pot veni și de la persoane fără cont. Această alegere arată legătura dintre cerințele aplicației și proiectarea bazei de date."
            ],
        ),
    ]

    for title, paragraphs in sections:
        add_heading(doc, title, 2)
        for paragraph in paragraphs:
            add_para(doc, paragraph)


def main():
    doc = Document()
    configure_doc(doc)
    add_footer(doc.sections[0])

    cover(doc, title_page=False)
    doc.add_page_break()
    cover(doc, title_page=True)
    doc.add_page_break()
    abstracts(doc)
    doc.add_page_break()
    front_matter(doc)
    doc.add_page_break()
    introduction(doc)
    doc.add_page_break()
    requirements(doc)
    doc.add_page_break()
    technologies(doc)
    doc.add_page_break()
    architecture(doc)
    doc.add_page_break()
    database_api(doc)
    doc.add_page_break()
    implementation(doc)
    doc.add_page_break()
    testing(doc)
    doc.add_page_break()
    conclusions(doc)
    doc.add_page_break()
    bibliography(doc)
    doc.add_page_break()
    appendices(doc)
    doc.add_page_break()
    code_screenshots_appendix(doc)
    doc.add_page_break()
    expand_content(doc)

    doc.core_properties.title = "Dezvoltarea unei aplicații web pentru gestionarea unui cabinet veterinar - VetCare"
    doc.core_properties.subject = "Lucrare de licență"
    doc.core_properties.author = "[Numele studentului]"
    doc.core_properties.keywords = "VetCare, Angular, Node.js, Express, MySQL, JWT, cabinet veterinar"
    doc.save(OUT)


if __name__ == "__main__":
    main()
